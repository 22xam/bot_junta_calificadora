import json
import sys
from docx import Document
import re
import os

def cargar_json(path):
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"❌ Error al cargar JSON: {e}")
        sys.exit(1)

def detectar_campos(doc):
    texto = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texto += "\n" + cell.text
    return set(re.findall(r'_\w+', texto))

def preparar_datos(json_data, campos_docx):
    datos = {
        "_MATERIA": json_data.get("materia", "")
    }

    for campo in campos_docx:
        if campo not in datos:
            datos[campo] = input(f"🖋️ Ingrese valor para {campo}: ").strip()

    return datos

def reemplazar_campos(doc, datos):
    for p in doc.paragraphs:
        for key, val in datos.items():
            if key in p.text:
                p.text = p.text.replace(key, val)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in datos.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, val)

def completar_listado(doc, listado):
    tablas_validas = []

    for table in doc.tables:
        if len(table.rows) > 1 and len(table.rows[0].cells) >= 6:
            encabezados = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if 'apellido' in encabezados[2].lower() and 'decisión' in encabezados[3].lower():
                tablas_validas.append(table)

    idx_reg = 0
    total = len(listado)

    for tabla in tablas_validas:
        filas_totales = len(tabla.rows)
        filas_datos = filas_totales - 1  # restamos 1 fila de encabezado
        registros_restantes = total - idx_reg
        filas_a_usar = min(filas_datos, registros_restantes)

        for i in range(1, 1 + filas_a_usar):
            fila = tabla.rows[i]
            fila.cells[0].text = str(listado[idx_reg].get("Orden", ""))
            fila.cells[1].text = str(listado[idx_reg].get("Puntos", ""))
            fila.cells[2].text = str(listado[idx_reg].get("Apellido y Nombre", ""))
            fila.cells[3].text = ""
            if len(fila.cells) > 4:
                fila.cells[4].text = ""
            if len(fila.cells) > 5:
                fila.cells[5].text = ""
            idx_reg += 1

        if idx_reg >= total:
            break  # ya completamos todos

    if idx_reg < total:
        print(f"⚠️ Atención: quedaron {total - idx_reg} registros sin colocar. Verificá si necesitás agregar más planillas.")

    tablas_validas = []

    # Detectar tablas con al menos 6 columnas y encabezado de notificación
    for table in doc.tables:
        if len(table.rows) > 1 and len(table.rows[0].cells) >= 6:
            encabezados = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if 'apellido' in encabezados[2].lower() and 'decisión' in encabezados[3].lower():
                tablas_validas.append(table)

    idx_reg = 0  # Índice global de registros

    for tabla in tablas_validas:
        filas_disponibles = len(tabla.rows) - 1  # quitamos encabezado
        for i in range(1, len(tabla.rows)):
            if idx_reg >= len(listado):
                return  # todos los registros insertados
            fila = tabla.rows[i]
            fila.cells[0].text = str(listado[idx_reg].get("Orden", ""))
            fila.cells[1].text = str(listado[idx_reg].get("Puntos", ""))
            fila.cells[2].text = str(listado[idx_reg].get("Apellido y Nombre", ""))
            fila.cells[3].text = ""
            if len(fila.cells) > 4:
                fila.cells[4].text = ""
            if len(fila.cells) > 5:
                fila.cells[5].text = ""
            idx_reg += 1  # avanzar al siguiente registro

def main():
    if len(sys.argv) != 2:
        print("🔧 Uso: python llenar_docx.py datos.json")
        sys.exit(1)

    json_path = sys.argv[1]
    base_docx = "BASE.docx"
    salida_docx = "salida.docx"

    if not os.path.exists(base_docx):
        print("❌ No se encontró BASE.docx en la carpeta actual.")
        sys.exit(1)

    print("📥 Cargando JSON...")
    json_data = cargar_json(json_path)

    print("📄 Abriendo documento...")
    doc = Document(base_docx)

    print("🔍 Detectando campos...")
    campos = detectar_campos(doc)

    print("🧠 Preparando datos...")
    datos = preparar_datos(json_data, campos)

    print("✍️ Reemplazando texto...")
    reemplazar_campos(doc, datos)

    print("📋 Rellenando tabla de puntajes...")
    completar_listado(doc, json_data.get("listado", []))

    doc.save(salida_docx)
    print(f"✅ Documento generado correctamente: {salida_docx}")

if __name__ == "__main__":
    main()
